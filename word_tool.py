from __future__ import annotations

import argparse
import os
import sys
from pathlib import Path
from typing import Iterable

from docx import Document


def resolve_path(path_text: str) -> Path:
    return Path(path_text).expanduser().resolve()


def ensure_docx_path(path: Path) -> None:
    if path.suffix.lower() != ".docx":
        raise ValueError("当前工具的编辑功能只支持 .docx 文件。")


def iter_paragraphs(container) -> Iterable:
    if hasattr(container, "paragraphs"):
        for paragraph in container.paragraphs:
            yield paragraph
    if hasattr(container, "tables"):
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_paragraphs(cell)
    if hasattr(container, "sections"):
        for section in container.sections:
            yield from iter_paragraphs(section.header)
            yield from iter_paragraphs(section.footer)


def open_in_default_app(path: Path) -> None:
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")
    os.startfile(path)
    print(f"已打开: {path}")


def create_document(path: Path, title: str | None, paragraphs: list[str]) -> None:
    ensure_docx_path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    if title:
        document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"已创建文档: {path}")


def append_paragraphs(path: Path, paragraphs: list[str]) -> None:
    ensure_docx_path(path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")

    document = Document(path)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"已追加 {len(paragraphs)} 段内容到: {path}")


def replace_text(path: Path, old_text: str, new_text: str) -> int:
    ensure_docx_path(path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")
    if not old_text:
        raise ValueError("--old 不能为空。")

    document = Document(path)
    replaced_count = 0

    for paragraph in iter_paragraphs(document):
        current_text = paragraph.text
        if old_text not in current_text:
            continue
        replaced_count += current_text.count(old_text)
        # 直接替换整段文本，能覆盖跨 run 的情况，但该段原有局部样式会被重建。
        paragraph.text = current_text.replace(old_text, new_text)

    document.save(path)
    return replaced_count


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="打开并编辑 Word 文档（编辑功能支持 .docx）。"
    )
    subparsers = parser.add_subparsers(dest="command", required=True)

    open_parser = subparsers.add_parser("open", help="用系统默认程序打开文档。")
    open_parser.add_argument("path", help="Word 文档路径。")

    create_parser = subparsers.add_parser("create", help="创建一个新的 .docx 文档。")
    create_parser.add_argument("path", help="目标 .docx 路径。")
    create_parser.add_argument("--title", help="文档标题。")
    create_parser.add_argument(
        "--text",
        action="append",
        default=[],
        help="文档段落内容。可重复传入多次。",
    )
    create_parser.add_argument(
        "--open-after",
        action="store_true",
        help="创建后立即打开文档。",
    )

    append_parser = subparsers.add_parser("append", help="向现有 .docx 末尾追加段落。")
    append_parser.add_argument("path", help="目标 .docx 路径。")
    append_parser.add_argument(
        "--text",
        action="append",
        required=True,
        help="要追加的段落内容。可重复传入多次。",
    )
    append_parser.add_argument(
        "--open-after",
        action="store_true",
        help="追加后立即打开文档。",
    )

    replace_parser = subparsers.add_parser("replace", help="替换 .docx 中的文本。")
    replace_parser.add_argument("path", help="目标 .docx 路径。")
    replace_parser.add_argument("--old", required=True, help="要被替换的文本。")
    replace_parser.add_argument("--new", required=True, help="替换成的新文本。")
    replace_parser.add_argument(
        "--open-after",
        action="store_true",
        help="替换后立即打开文档。",
    )

    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()

    try:
        path = resolve_path(args.path)

        if args.command == "open":
            open_in_default_app(path)
            return 0

        if args.command == "create":
            create_document(path, args.title, args.text)
            if args.open_after:
                open_in_default_app(path)
            return 0

        if args.command == "append":
            append_paragraphs(path, args.text)
            if args.open_after:
                open_in_default_app(path)
            return 0

        if args.command == "replace":
            replaced_count = replace_text(path, args.old, args.new)
            print(f"已完成替换，匹配次数: {replaced_count}")
            if args.open_after:
                open_in_default_app(path)
            return 0
    except Exception as exc:  # pragma: no cover
        print(f"执行失败: {exc}", file=sys.stderr)
        return 1

    parser.print_help()
    return 1


if __name__ == "__main__":
    raise SystemExit(main())

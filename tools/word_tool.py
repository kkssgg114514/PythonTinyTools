from __future__ import annotations

import argparse
import sys
from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

TITLE_FONT_NAME = "方正小标宋简体"
BODY_FONT_NAME = "仿宋_GB2312"
TITLE_FONT_SIZE = Pt(22)
BODY_FONT_SIZE = Pt(16)


def resolve_path(path_text: str) -> Path:
    return Path(path_text).expanduser().resolve()


def ensure_docx_file(path: Path) -> None:
    if path.suffix.lower() != ".docx":
        raise ValueError("当前脚本仅支持 .docx 文件。")
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")


def iter_paragraphs(container) -> Iterable:
    if hasattr(container, "paragraphs"):
        for paragraph in container.paragraphs:
            yield paragraph
    if hasattr(container, "tables"):
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_paragraphs(cell)


def find_title_paragraph(document: Document):
    for paragraph in iter_paragraphs(document):
        if paragraph.text.strip():
            return paragraph
    return None


def apply_run_style(run, font_name: str, font_size: Pt) -> None:
    run.font.name = font_name
    run.font.size = font_size
    r_fonts = run._element.rPr.rFonts
    for font_key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(font_key), font_name)


def apply_paragraph_style(paragraph, font_name: str, font_size: Pt) -> None:
    if not paragraph.text.strip():
        return
    for run in paragraph.runs:
        apply_run_style(run, font_name, font_size)


def standardize_document(path: Path) -> tuple[int, int]:
    ensure_docx_file(path)

    document = Document(path)
    title_paragraph = find_title_paragraph(document)
    if title_paragraph is None:
        raise ValueError("文档中没有可处理的非空段落。")

    title_element = title_paragraph._element
    title_count = 0
    body_count = 0

    for paragraph in iter_paragraphs(document):
        if not paragraph.text.strip():
            continue
        if paragraph._element is title_element:
            apply_paragraph_style(paragraph, TITLE_FONT_NAME, TITLE_FONT_SIZE)
            title_count += 1
            continue
        apply_paragraph_style(paragraph, BODY_FONT_NAME, BODY_FONT_SIZE)
        body_count += 1

    document.save(path)
    return title_count, body_count


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="标准化 Word .docx 文档格式，并覆盖原文件。"
    )
    parser.add_argument(
        "path",
        help="要处理的 Word .docx 文件路径，支持绝对路径或相对路径。",
    )
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()

    try:
        path = resolve_path(args.path)
        title_count, body_count = standardize_document(path)
        print(f"处理完成: {path}")
        print(f"标题段落: {title_count}")
        print(f"正文段落: {body_count}")
        return 0
    except Exception as exc:  # pragma: no cover
        print(f"执行失败: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
